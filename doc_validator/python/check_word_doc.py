"""Проверка и автоисправление оформления Word-документа по «Памятке».

Использует python-docx (читает .docx как ZIP, Word не нужен).

    python check_word_doc.py "путь/к/файлу.docx"            # только проверить
    python check_word_doc.py "путь/к/файлу.docx" --fix       # спросит и исправит, сохранит в *_fixed.docx
    python check_word_doc.py "путь/к/файлу.docx" --fix --yes # без вопросов
    python check_word_doc.py "путь/к/файлу.docx" --fix --in-place  # переписать исходник (с .bak)
"""
from __future__ import annotations

import argparse
import re
import shutil
import sys
from dataclasses import dataclass
from pathlib import Path
from typing import List

from docx import Document
from docx.enum.text import WD_LINE_SPACING, WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt
from docx.text.paragraph import Paragraph
from lxml import etree

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

# ---------- регулярки для шапки ----------
RE_FIO_DOT       = re.compile(r"\b[А-ЯЁ][а-яё]+\s+[А-ЯЁ]\.\s?[А-ЯЁ]\.")
RE_FULL_NAME     = re.compile(r"\b[А-ЯЁ][а-яё]{2,}\s+[А-ЯЁ][а-яё]{2,}\s+[А-ЯЁ][а-яё]{2,}\b")
RE_FIO_BAD_GLUED = re.compile(r"[А-ЯЁ][а-яё]+[А-ЯЁ]\.")
RE_FIO_ONE_INIT  = re.compile(r"\b[А-ЯЁ][а-яё]+\s+[А-ЯЁ]\.(?!\s?[А-ЯЁ]\.)")
RE_TITLE_O       = re.compile(r"^\s*Об?\s+[а-яёА-ЯЁ]")
# Корни наименований госорганов/орг-форм без жёсткой привязки к падежу
RE_ORG_NOM       = re.compile(r"\b(ООО|АО|ПАО|ОАО|Министерств|Управлени|Федеральн|Правительств|Администраци|Канцеляри|ФНС|УФНС|ИФНС|Межрегиональн|Инспекци|Департамент|Комитет|Служб)")

# ---------- параметры из Памятки ----------
MARGIN_LEFT_MM   = 30
MARGIN_TOP_MM    = 20
MARGIN_BOTTOM_MM = 20
MARGIN_RIGHT_MM  = 10
MM_TOL           = 0.5
CM_TOL           = 0.05
FONT_REQ         = "Times New Roman"
INDENT_REQ_CM    = 1.25
SIZES_OK         = {13, 14}
FIX_FONT_SIZE_PT = 14
FIX_LINE_SPACING = 1.5

EMU_PER_MM = 36000
EMU_PER_CM = 360000


@dataclass
class Issue:
    code: str
    text: str
    fixable: bool


def emu_to_mm(v) -> float:
    return 0.0 if v is None else round(int(v) / EMU_PER_MM, 2)


def emu_to_cm(v) -> float:
    return 0.0 if v is None else round(int(v) / EMU_PER_CM, 3)


def _run_font(run) -> str | None:
    name = run.font.name
    if name:
        return name
    try:
        return run.style.font.name
    except Exception:
        return None


def _is_body_para(p) -> bool:
    text = p.text.strip()
    if len(text) <= 80:
        return False
    style_name = (p.style.name or "").lower()
    return "heading" not in style_name and "заголовок" not in style_name


def _top_paragraphs(doc, limit: int = 30) -> List[Paragraph]:
    """Параграфы в документном порядке, включая ячейки таблиц (для шапки).
    Итерация идёт по сырому XML — это снимает проблему дублирования
    параграфов в объединённых ячейках, которое даёт Table.rows[*].cells."""
    result: List[Paragraph] = []
    body = doc.element.body
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            result.append(Paragraph(child, doc))
        elif child.tag == qn("w:tbl"):
            for tr in child.iter(qn("w:tr")):
                for tc in tr.findall(qn("w:tc")):
                    for p_elem in tc.findall(qn("w:p")):
                        result.append(Paragraph(p_elem, doc))
        if len(result) >= limit:
            break
    return result[:limit]


def _left_indent_cm(p: Paragraph) -> float:
    li = p.paragraph_format.left_indent
    return emu_to_cm(li) if li is not None else 0.0


def _in_right_block(p: Paragraph) -> bool:
    """Эвристика: параграф относится к правому блоку шапки.
    True если: выравнивание по правому краю, или левый отступ > 7 см,
    или параграф находится в последней w:tc своей строки таблицы."""
    if p.alignment == WD_ALIGN_PARAGRAPH.RIGHT:
        return True
    if _left_indent_cm(p) > 7.0:
        return True
    tc = p._element.getparent()
    if tc is None or tc.tag != qn("w:tc"):
        return False
    tr = tc.getparent()
    if tr is None or tr.tag != qn("w:tr"):
        return False
    tcs = tr.findall(qn("w:tc"))
    return len(tcs) >= 2 and tcs[-1] is tc


# ---------- проверки ----------
def _has_letterhead(doc) -> bool:
    """True, если документ начинается с таблицы (типичная разметка бланка с
    гербом / реквизитами организации). В таких документах верхнее поле обычно
    уменьшено, чтобы герб уместился — это нормально и проверяться не должно."""
    body = doc.element.body
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            text = "".join((t.text or "") for t in child.iter(qn("w:t"))).strip()
            if text:
                return False
            # пустой параграф в начале — пропускаем
            continue
        if child.tag == qn("w:tbl"):
            return True
        if child.tag == qn("w:sectPr"):
            continue
        # любой иной первый значимый элемент — это уже не бланк
        break
    return False


def collect_issues(doc) -> List[Issue]:
    issues: List[Issue] = []
    s = doc.sections[0]
    has_letterhead = _has_letterhead(doc)
    for name, actual, required, code in (
        ("Левое",   emu_to_mm(s.left_margin),   MARGIN_LEFT_MM,   "MARGIN_L"),
        ("Верхнее", emu_to_mm(s.top_margin),    MARGIN_TOP_MM,    "MARGIN_T"),
        ("Нижнее",  emu_to_mm(s.bottom_margin), MARGIN_BOTTOM_MM, "MARGIN_B"),
        ("Правое",  emu_to_mm(s.right_margin),  MARGIN_RIGHT_MM,  "MARGIN_R"),
    ):
        if code == "MARGIN_T" and has_letterhead:
            # Бланк с гербом «съедает» верхнее поле — это интенционально.
            continue
        if abs(actual - required) > MM_TOL:
            issues.append(Issue(code, f"{name} поле {actual} мм (нужно {required} мм).", True))

    bad_fonts: set[str] = set()
    bad_sizes: set[str] = set()
    bad_spacing = False
    body_paras = 0
    bad_indent = 0

    for p in doc.paragraphs:
        text = p.text.strip()
        if len(text) <= 1:
            continue
        for run in p.runs:
            name = _run_font(run)
            if name and name != FONT_REQ:
                bad_fonts.add(name)
            if run.font.size is not None and run.font.size.pt not in SIZES_OK:
                bad_sizes.add(f"{run.font.size.pt:g}")

        pf = p.paragraph_format
        rule, ls = pf.line_spacing_rule, pf.line_spacing
        ok = False
        if rule in (WD_LINE_SPACING.SINGLE, WD_LINE_SPACING.ONE_POINT_FIVE):
            ok = True
        elif rule == WD_LINE_SPACING.MULTIPLE and ls is not None:
            ok = 0.95 <= float(ls) <= 1.55
        elif rule is None and ls is None:
            ok = True
        elif ls is not None and 0.95 <= float(ls) <= 1.55:
            ok = True
        if not ok:
            bad_spacing = True

        if _is_body_para(p):
            body_paras += 1
            ind = pf.first_line_indent
            ind_cm = emu_to_cm(ind) if ind is not None else 0.0
            if abs(ind_cm - INDENT_REQ_CM) > CM_TOL:
                bad_indent += 1

    if bad_fonts:
        issues.append(Issue("FONT",   f"Шрифт не {FONT_REQ}: {', '.join(sorted(bad_fonts))}.", True))
    if bad_sizes:
        issues.append(Issue("SIZE",   f"Размер шрифта вне 13–14 пт: {', '.join(sorted(bad_sizes))}.", True))
    if bad_spacing:
        issues.append(Issue("LINESP", "Межстрочный интервал вне 1.0–1.5.", True))
    if body_paras and bad_indent:
        issues.append(Issue("INDENT", f"Абзацный отступ ≠ 1.25 см: {bad_indent} из {body_paras} абзацев.", True))

    # номер страницы — эвристика
    if len(doc.paragraphs) >= 40:
        found_page = False
        centered = False
        decor = False
        for section in doc.sections:
            hdr = section.header
            xml = hdr._element.xml if hdr is not None else ""
            if "PAGE" in xml and "instr" in xml.lower():
                found_page = True
            for hp in hdr.paragraphs:
                if hp.alignment == WD_ALIGN_PARAGRAPH.CENTER:
                    centered = True
                t = hp.text
                if "-" in t or "стр" in t.lower() or "с." in t.lower():
                    decor = True
        if not found_page:
            issues.append(Issue("PAGE_ADD", "Похоже на многостраничный документ, но номер страницы не найден.", True))
        else:
            if not centered:
                issues.append(Issue("PAGE_ALIGN", "Номер страницы не по центру верхнего поля.", True))
            if decor:
                issues.append(Issue("PAGE_DECOR", "В номере страницы лишние символы.", True))

    full = "\n".join(p.text for p in doc.paragraphs)
    if not re.search(r"8\(\d{3,5}\)\s?\d", full):
        issues.append(Issue("EXEC", "Не найдена отметка об исполнителе (телефон 8(код)номер).", False))

    _check_header(doc, issues)
    _check_blank(doc, issues, _blank_template_path(getattr(collect_issues, "_blank_override", None)))

    return issues


def _blank_template_path(override: str | None = None) -> Path:
    """Путь к blank_template.txt: --blank-template > рядом со скриптом > в корне doc_validator."""
    if override:
        return Path(override)
    here = Path(__file__).resolve().parent
    candidates = [here / "blank_template.txt", here.parent / "blank_template.txt"]
    for c in candidates:
        if c.exists():
            return c
    return candidates[-1]


def _load_blank_template(config_path: Path) -> list:
    """Парсит blank_template.txt → список {pattern, is_regex, optional, raw}."""
    if not config_path.exists():
        return []
    items = []
    for raw in config_path.read_text(encoding="utf-8").splitlines():
        line = raw.strip()
        if not line or line.startswith("#"):
            continue
        optional = False
        is_regex = False
        if line.startswith("?"):
            optional = True
            line = line[1:].lstrip()
        if line.startswith("re:"):
            is_regex = True
            line = line[3:]
        if not line:
            continue
        items.append({
            "pattern": line,
            "is_regex": is_regex,
            "optional": optional,
            "raw": raw.strip(),
        })
    return items


def _blank_text(doc) -> str:
    """Текст левой части шапки. Берёт первую таблицу (типичный бланк-таблица)
    и собирает уникальные w:tc в левой половине; если таблицы нет — собирает
    параграфы вне правого блока в верхней части документа."""
    if doc.tables:
        tbl = doc.tables[0]._element
        all_rows = tbl.findall(qn("w:tr"))
        # Найти максимальное число tc по строкам — это «ширина» сетки в tc
        max_tcs = max((len(tr.findall(qn("w:tc"))) for tr in all_rows), default=0)
        half = max(1, max_tcs // 2) if max_tcs >= 2 else 1
        seen = set()
        parts = []
        for tr in all_rows:
            tcs = tr.findall(qn("w:tc"))
            for idx, tc in enumerate(tcs):
                if idx >= half:
                    break  # это уже правая половина
                if id(tc) in seen:
                    continue
                seen.add(id(tc))
                for p_elem in tc.findall(qn("w:p")):
                    text = Paragraph(p_elem, doc).text
                    if text.strip():
                        parts.append(text)
        return "\n".join(parts)
    # без таблицы — параграфы из шапки, которые НЕ в правом блоке
    top = _top_paragraphs(doc, limit=60)
    return "\n".join(p.text for p in top if not _in_right_block(p) and p.text.strip())


def _check_blank(doc, issues: List[Issue], config_path: Path) -> None:
    items = _load_blank_template(config_path)
    if not items:
        # Конфиг не настроен — молча пропускаем.
        # Чтобы пользователь видел, что проверка отключена, можно добавить
        # информационное сообщение, но я оставляю «тихо».
        return
    blank_text = _blank_text(doc)
    for item in items:
        if item["is_regex"]:
            try:
                found = bool(re.search(item["pattern"], blank_text))
            except re.error:
                issues.append(Issue("BLANK_BAD_REGEX",
                    f"Некорректное regexp в blank_template.txt: «{item['raw']}».", False))
                continue
        else:
            found = item["pattern"] in blank_text
        if not found:
            label = item["pattern"] if not item["is_regex"] else f"(regex) {item['pattern']}"
            if item["optional"]:
                issues.append(Issue("BLANK_OPTIONAL",
                    f"В бланке не найдено (необязательно): «{label}».", False))
            else:
                issues.append(Issue("BLANK_MISSING",
                    f"В бланке не найдено: «{label}».", False))


def _check_header(doc, issues: List[Issue]) -> None:
    """Проверка шапки документа (правая часть): гриф ДСП, адресат, инициалы, заголовок."""
    top = _top_paragraphs(doc, limit=60)
    if not top:
        return

    right_block = [p for p in top if _in_right_block(p) and p.text.strip()]
    right_text = "\n".join(p.text for p in right_block)
    full_top_text = "\n".join(p.text for p in top)

    # Группируем параграфы правой части по родительской ячейке (w:tc),
    # чтобы отличать «ссылку на № входящего» от собственно адресата.
    from collections import OrderedDict
    groups: "OrderedDict[int, list]" = OrderedDict()
    for p in right_block:
        tc = p._element.getparent()
        key = id(tc) if tc is not None and tc.tag == qn("w:tc") else 0
        groups.setdefault(key, []).append(p)

    adresat_paras: List[Paragraph] = []
    for paras in groups.values():
        joined = "\n".join(p.text for p in paras)
        if RE_FIO_DOT.search(joined) or RE_FULL_NAME.search(joined) or RE_ORG_NOM.search(joined):
            adresat_paras = paras
            break

    # --- Гриф ДСП ---
    dsp_in_doc = bool(re.search(r"(?i)для служебного пользования", full_top_text))
    if dsp_in_doc:
        dsp_in_right = bool(re.search(r"(?i)для служебного пользования", right_text))
        if not dsp_in_right:
            issues.append(Issue(
                "DSP_POS",
                "«Для служебного пользования» найдено, но не в правом верхнем углу.",
                False,
            ))
        if not re.search(r"(?i)экз\.?\s*№\s*\d", full_top_text):
            issues.append(Issue(
                "DSP",
                "«Для служебного пользования» без номера экземпляра («Экз. № …»).",
                False,
            ))

    # --- Адресат ---
    if not right_block:
        issues.append(Issue(
            "ADRESAT_MISSING",
            "Не найден блок адресата в правой верхней части документа.",
            False,
        ))
    elif not adresat_paras:
        issues.append(Issue(
            "ADRESAT_MISSING",
            "Блок справа есть, но в нём не найдены ни ФИО, ни наименование организации.",
            False,
        ))
    else:
        # Проверка выравнивания — только если адресат не в таблице
        # (в табличном бланке выравнивание задаёт сама ячейка).
        in_table = any(
            p._element.getparent() is not None
            and p._element.getparent().tag == qn("w:tc")
            for p in adresat_paras
        )
        if not in_table and len(adresat_paras) >= 2:
            alignments = {p.alignment for p in adresat_paras}
            indents    = {round(_left_indent_cm(p), 1) for p in adresat_paras}
            if len(alignments) > 1 or len(indents) > 1:
                issues.append(Issue(
                    "ADRESAT_ALIGN",
                    "Строки адресата выровнены неодинаково (должны быть по левому краю блока).",
                    False,
                ))

        for p in adresat_paras:
            pf = p.paragraph_format
            rule, ls = pf.line_spacing_rule, pf.line_spacing
            bad = False
            if rule == WD_LINE_SPACING.MULTIPLE and ls is not None and float(ls) > 1.05:
                bad = True
            elif rule in (WD_LINE_SPACING.ONE_POINT_FIVE, WD_LINE_SPACING.DOUBLE):
                bad = True
            if bad:
                issues.append(Issue(
                    "ADRESAT_SPACING",
                    "Межстрочный интервал в адресате не 1.0.",
                    False,
                ))
                break

    # --- Формат инициалов ---
    if RE_FIO_BAD_GLUED.search(right_text):
        issues.append(Issue(
            "INITIALS_BAD",
            "Инициалы записаны слитно с фамилией (нужно «Фамилия И.О.» через пробел).",
            False,
        ))
    elif RE_FIO_ONE_INIT.search(right_text) and not RE_FIO_DOT.search(right_text):
        issues.append(Issue(
            "INITIALS_BAD",
            "У фамилии указан только один инициал (нужно два: «Фамилия И.О.»).",
            False,
        ))

    # --- Заголовок «О...» / «Об...» ---
    # Ищем в пределах всей шапки (порядок обхода ячеек таблицы может не совпадать
    # с визуальным «под адресатом»).
    title_found = any(
        RE_TITLE_O.match(p.text.strip())
        for p in top
        if p.text.strip()
    )
    if not title_found:
        issues.append(Issue(
            "TITLE_MISSING",
            "Не найден заголовок к тексту «О …»/«Об …» в шапке (п. 4 Памятки).",
            False,
        ))


# ---------- исправления ----------
def fix_margins(doc, codes: set[str]) -> None:
    s = doc.sections[0]
    if "MARGIN_L" in codes: s.left_margin   = Mm(MARGIN_LEFT_MM)
    if "MARGIN_T" in codes: s.top_margin    = Mm(MARGIN_TOP_MM)
    if "MARGIN_B" in codes: s.bottom_margin = Mm(MARGIN_BOTTOM_MM)
    if "MARGIN_R" in codes: s.right_margin  = Mm(MARGIN_RIGHT_MM)


def _set_rpr_font_name(run, name: str) -> None:
    # Гарантированно ставим имя шрифта для всех скриптов (latin, cs, eastAsia)
    run.font.name = name
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = etree.SubElement(rPr, qn("w:rFonts"))
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rFonts.set(qn(attr), name)


def fix_fonts_sizes_spacing_indent(doc, codes: set[str]) -> None:
    for p in doc.paragraphs:
        text = p.text.strip()
        if len(text) <= 1:
            continue
        for run in p.runs:
            if "FONT" in codes:
                name = _run_font(run)
                if not name or name != FONT_REQ:
                    _set_rpr_font_name(run, FONT_REQ)
            if "SIZE" in codes:
                sz = run.font.size
                if sz is None or sz.pt not in SIZES_OK:
                    run.font.size = Pt(FIX_FONT_SIZE_PT)
        if "LINESP" in codes:
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            p.paragraph_format.line_spacing      = FIX_LINE_SPACING
        if "INDENT" in codes and _is_body_para(p):
            p.paragraph_format.first_line_indent = Cm(INDENT_REQ_CM)


def fix_page_numbers(doc, codes: set[str]) -> None:
    """Записывает в primary header абзац с полем PAGE по центру."""
    if not (codes & {"PAGE_ADD", "PAGE_ALIGN", "PAGE_DECOR"}):
        return
    W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    for section in doc.sections:
        hdr = section.header
        # очистить
        for p in list(hdr.paragraphs):
            p._element.getparent().remove(p._element)
        p = hdr.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        r = run._element
        # <w:fldChar begin/> <w:instrText>PAGE</w:instrText> <w:fldChar end/>
        fld_begin = etree.SubElement(r, qn("w:fldChar"))
        fld_begin.set(qn("w:fldCharType"), "begin")
        instr = etree.SubElement(r, qn("w:instrText"))
        instr.set(qn("xml:space"), "preserve")
        instr.text = "PAGE"
        fld_end = etree.SubElement(r, qn("w:fldChar"))
        fld_end.set(qn("w:fldCharType"), "end")


def apply_fixes(doc, issues: List[Issue]) -> None:
    codes = {i.code for i in issues if i.fixable}
    fix_margins(doc, codes)
    fix_fonts_sizes_spacing_indent(doc, codes)
    fix_page_numbers(doc, codes)


# ---------- вывод ----------
def print_issues(issues: List[Issue]) -> None:
    if not issues:
        print("OK. Все автоматические проверки пройдены.")
        return
    for i, iss in enumerate(issues, 1):
        tag = "[можно исправить]" if iss.fixable else "[ручная правка] "
        print(f"{i}. {tag} {iss.text}")
    fixable = sum(1 for i in issues if i.fixable)
    print(f"\nАвтоисправимых: {fixable}; требуют ручной правки: {len(issues) - fixable}")


# ---------- main ----------
def parse_args() -> argparse.Namespace:
    ap = argparse.ArgumentParser()
    ap.add_argument("path", nargs="?", help="Путь к .docx")
    ap.add_argument("--fix", action="store_true", help="Применить автоисправление")
    ap.add_argument("--yes", "-y", action="store_true", help="Не задавать вопрос «исправить?»")
    ap.add_argument("--in-place", action="store_true", help="Перезаписать исходный файл (с .bak)")
    ap.add_argument("--blank-template", help="Путь к blank_template.txt (по умолчанию — рядом со скриптом)")
    return ap.parse_args()


def main() -> int:
    args = parse_args()
    if args.path:
        path = Path(args.path)
    else:
        candidates = sorted(Path.cwd().glob("*.docx"))
        if not candidates:
            print("Укажите путь к .docx или положите файл в текущую папку.")
            return 2
        path = candidates[0]
    if not path.exists():
        print(f"Файл не найден: {path}")
        return 2

    if args.blank_template:
        collect_issues._blank_override = args.blank_template  # type: ignore[attr-defined]
    doc = Document(str(path))
    issues = collect_issues(doc)

    print(f"\nДокумент: {path}")
    print("-" * 60)
    print_issues(issues)

    if not args.fix or not issues:
        return 0 if not issues else 1

    fixable = [i for i in issues if i.fixable]
    if not fixable:
        print("\nАвтоисправление неприменимо — все замечания требуют ручной правки.")
        return 1

    if not args.yes:
        ans = input(f"\nИсправить {len(fixable)} автоисправимых замечаний? [Y/N]: ").strip().lower()
        if ans not in {"y", "yes", "д", "да"}:
            return 1

    apply_fixes(doc, issues)

    if args.in_place:
        shutil.copy2(path, str(path) + ".bak")
        doc.save(str(path))
        out = path
    else:
        out = path.with_name(f"{path.stem}_fixed.docx")
        doc.save(str(out))
    print(f"\nСохранено: {out}")

    issues2 = collect_issues(Document(str(out)))
    print("\nПерепроверка:")
    print("-" * 60)
    print_issues(issues2)
    return 0 if not issues2 else 1


if __name__ == "__main__":
    sys.exit(main())
