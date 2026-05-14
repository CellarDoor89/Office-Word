"""DocCheck GUI — drag-and-drop проверка оформления документа по «Памятке».

Запуск из исходников:
    pip install python-docx tkinterdnd2
    python app.py

Сборка одного .exe:
    build.bat     (нужен Python 3.10+ на машине сборки)

В готовом .exe ничего ставить не надо — все зависимости внутри.
"""
from __future__ import annotations

import os
import sys
import threading
import traceback
from pathlib import Path

import tkinter as tk
from tkinter import filedialog, messagebox, ttk, scrolledtext

try:
    from tkinterdnd2 import DND_FILES, TkinterDnD
    HAS_DND = True
except Exception:
    HAS_DND = False


def _bundle_dir() -> Path:
    """Папка, где лежат данные программы.
    В режиме PyInstaller — sys._MEIPASS (временная распаковка onefile).
    В обычном запуске — папка скрипта."""
    return Path(getattr(sys, "_MEIPASS", Path(__file__).resolve().parent))


def _exe_dir() -> Path:
    """Папка рядом с .exe (для PyInstaller) или со скриптом."""
    if getattr(sys, "frozen", False):
        return Path(sys.executable).resolve().parent
    return Path(__file__).resolve().parent


def _setup_imports() -> None:
    """Найти check_word_doc.py и добавить его папку в sys.path."""
    candidates = [
        _bundle_dir(),
        _bundle_dir() / "python",
        _exe_dir().parent / "python",
        _exe_dir() / "python",
    ]
    for c in candidates:
        if (c / "check_word_doc.py").exists():
            sys.path.insert(0, str(c))
            return
    raise ImportError("Не нашёл check_word_doc.py — обратитесь к разработчику.")


_setup_imports()
import check_word_doc as v  # noqa: E402
from docx import Document   # noqa: E402


def _ensure_blank_template_external() -> Path | None:
    """Возвращает путь к blank_template.txt, который пользователь может редактировать.

    В .exe (frozen): если файл рядом с .exe отсутствует — копируем дефолт
    из бандла, чтобы пользователю было что открыть в Блокноте.
    В dev-режиме: используем blank_template.txt из репозитория напрямую."""
    ext = _exe_dir() / "blank_template.txt"
    if ext.exists():
        return ext

    if getattr(sys, "frozen", False):
        bundled = _bundle_dir() / "blank_template.txt"
        if bundled.exists():
            try:
                ext.write_text(bundled.read_text(encoding="utf-8"), encoding="utf-8")
                return ext
            except OSError:
                return bundled
        return None

    # dev-режим — используем файл из репозитория, не копируя
    for candidate in (
        _exe_dir().parent / "blank_template.txt",
        _bundle_dir() / "blank_template.txt",
    ):
        if candidate.exists():
            return candidate
    return None


class App:
    def __init__(self, root: tk.Misc) -> None:
        self.root = root
        root.title("DocCheck — проверка оформления документа")
        root.geometry("820x640")
        root.minsize(640, 540)

        self.current_path: Path | None = None
        self.fixed_path: Path | None = None
        self.issues: list = []

        self._build_ui()
        self._init_blank_template_status()

    # ---------------------- UI ----------------------
    def _build_ui(self) -> None:
        style = ttk.Style()
        try:
            style.theme_use("vista")
        except tk.TclError:
            pass

        top = ttk.Frame(self.root, padding=12)
        top.pack(fill="x")

        if HAS_DND:
            hint = "⬇  Перетащите .docx-файл сюда\n\nили нажмите «Выбрать файл…»"
        else:
            hint = "Нажмите «Выбрать файл…», чтобы начать проверку"

        self.drop_zone = tk.Label(
            top, text=hint, height=6,
            bg="#eef3f8", fg="#1a3650",
            font=("Segoe UI", 12), relief="ridge", bd=2, cursor="hand2",
        )
        self.drop_zone.pack(fill="x", pady=(0, 8))
        self.drop_zone.bind("<Button-1>", lambda e: self._browse())

        if HAS_DND:
            self.drop_zone.drop_target_register(DND_FILES)
            self.drop_zone.dnd_bind("<<Drop>>", self._on_drop)

        btn_row = ttk.Frame(top)
        btn_row.pack(fill="x")
        ttk.Button(btn_row, text="Выбрать файл…", command=self._browse).pack(side="left")
        ttk.Button(btn_row, text="Открыть шаблон бланка…",
                   command=self._open_blank_template).pack(side="left", padx=8)

        self.file_label = ttk.Label(btn_row, text="(файл не выбран)", foreground="#666")
        self.file_label.pack(side="left", padx=12)

        # Результаты
        result_frame = ttk.LabelFrame(self.root, text="Результат проверки", padding=8)
        result_frame.pack(fill="both", expand=True, padx=12, pady=(0, 8))

        self.result_text = scrolledtext.ScrolledText(
            result_frame, height=14, wrap="word",
            font=("Consolas", 10), state="disabled",
        )
        self.result_text.pack(fill="both", expand=True)
        self.result_text.tag_config("ok",     foreground="#1a7a3a")
        self.result_text.tag_config("fix",    foreground="#b35900")
        self.result_text.tag_config("manual", foreground="#8a2a2a")
        self.result_text.tag_config("info",   foreground="#555")
        self.result_text.tag_config("bold",   font=("Consolas", 10, "bold"))

        # Кнопки действий
        actions = ttk.Frame(self.root, padding=(12, 0, 12, 12))
        actions.pack(fill="x")

        self.fix_btn = ttk.Button(
            actions, text="Исправить автоисправимые", command=self._fix, state="disabled",
        )
        self.fix_btn.pack(side="left")

        self.open_fixed_btn = ttk.Button(
            actions, text="Открыть исправленный файл",
            command=self._open_fixed, state="disabled",
        )
        self.open_fixed_btn.pack(side="left", padx=8)

        ttk.Button(actions, text="Очистить", command=self._clear).pack(side="right")

        # Статусная строка про шаблон бланка
        self.status = ttk.Label(self.root, text="", foreground="#666", padding=(12, 0, 12, 8))
        self.status.pack(fill="x")

        self._set_text(
            "Перетащите файл в окно или выберите через кнопку.\n\n"
            "Программа проверит документ по «Памятке» — поля, шрифт, кегль, "
            "интервал, нумерацию страниц, реквизиты шапки и т.п.",
            tag="info",
        )

    def _init_blank_template_status(self) -> None:
        path = _ensure_blank_template_external()
        if path:
            self.status.config(text=f"Шаблон бланка: {path}")
        else:
            self.status.config(text="Шаблон бланка (blank_template.txt) не найден — проверка бланка отключена.")

    # ---------------------- File selection ----------------------
    def _on_drop(self, event) -> None:
        # event.data на Windows бывает в фигурных скобках, может содержать несколько путей
        raw = event.data.strip()
        # вытащить первый путь
        if raw.startswith("{") and "}" in raw:
            raw = raw[1:raw.index("}")]
        elif raw.startswith('"') and raw.endswith('"'):
            raw = raw[1:-1]
        else:
            raw = raw.split(" ")[0]
        self._load_file(Path(raw))

    def _browse(self) -> None:
        path = filedialog.askopenfilename(
            title="Выберите документ",
            filetypes=[("Word документы", "*.docx;*.doc"), ("Все файлы", "*.*")],
        )
        if path:
            self._load_file(Path(path))

    def _open_blank_template(self) -> None:
        p = _ensure_blank_template_external()
        if not p or not p.exists():
            messagebox.showwarning(
                "Шаблон не найден",
                "blank_template.txt не найден рядом с программой.",
            )
            return
        try:
            os.startfile(str(p))  # type: ignore[attr-defined]
        except AttributeError:
            # не-Windows
            import subprocess
            subprocess.Popen(["xdg-open", str(p)])

    def _load_file(self, path: Path) -> None:
        if not path.exists():
            messagebox.showerror("Ошибка", f"Файл не найден:\n{path}")
            return
        if path.suffix.lower() == ".doc":
            messagebox.showerror(
                "Старый формат .doc",
                "DocCheck умеет читать только .docx.\n\n"
                "Откройте файл в Word и сохраните как .docx "
                "(Файл → Сохранить как → Документ Word).",
            )
            return
        if path.suffix.lower() != ".docx":
            messagebox.showerror(
                "Неподдерживаемый формат",
                f"Ожидается .docx, а у вас «{path.suffix}».",
            )
            return
        self.current_path = path
        self.fixed_path = None
        self.file_label.config(text=path.name, foreground="black")
        self.open_fixed_btn.config(state="disabled")
        self._run_validate()

    # ---------------------- Validate / Fix ----------------------
    def _run_validate(self) -> None:
        self._set_text("Проверяю…\n", tag="info")
        self.fix_btn.config(state="disabled")

        def work() -> None:
            try:
                bt = _ensure_blank_template_external()
                if bt and bt.exists():
                    v.collect_issues._blank_override = str(bt)  # type: ignore[attr-defined]
                else:
                    if hasattr(v.collect_issues, "_blank_override"):
                        delattr(v.collect_issues, "_blank_override")
                doc = Document(str(self.current_path))
                issues = v.collect_issues(doc)
                self.root.after(0, lambda: self._show_issues(issues))
            except Exception:
                err = traceback.format_exc()
                self.root.after(0, lambda: self._show_error(err))

        threading.Thread(target=work, daemon=True).start()

    def _show_issues(self, issues: list) -> None:
        self.issues = issues
        self.result_text.config(state="normal")
        self.result_text.delete("1.0", "end")
        if not issues:
            self.result_text.insert(
                "end",
                "✓ Все автоматические проверки пройдены.\n\n"
                "Это не отменяет визуального осмотра перед отправкой на подпись.\n",
                "ok",
            )
            self.fix_btn.config(state="disabled")
        else:
            fixable = sum(1 for i in issues if i.fixable)
            manual = len(issues) - fixable
            for n, i in enumerate(issues, 1):
                tag = "fix" if i.fixable else "manual"
                marker = "[можно исправить] " if i.fixable else "[ручная правка]  "
                self.result_text.insert("end", f"{n:>2}. {marker}{i.text}\n", tag)
            self.result_text.insert(
                "end",
                f"\nАвтоисправимых: {fixable};  требуют ручной правки: {manual}\n",
                "info",
            )
            self.fix_btn.config(state="normal" if fixable else "disabled")
        self.result_text.config(state="disabled")

    def _show_error(self, err: str) -> None:
        self._set_text(f"Ошибка при проверке:\n\n{err}", tag="manual")

    def _fix(self) -> None:
        if not self.current_path or not self.issues:
            return
        self._set_text("Применяю исправления…\n", tag="info")
        self.fix_btn.config(state="disabled")

        def work() -> None:
            try:
                doc = Document(str(self.current_path))
                v.apply_fixes(doc, self.issues)
                out = self.current_path.with_name(
                    f"{self.current_path.stem}_fixed.docx"
                )
                doc.save(str(out))
                issues2 = v.collect_issues(Document(str(out)))
                self.fixed_path = out
                self.root.after(0, lambda: self._show_after_fix(out, issues2))
            except Exception:
                err = traceback.format_exc()
                self.root.after(0, lambda: self._show_error(err))

        threading.Thread(target=work, daemon=True).start()

    def _show_after_fix(self, out: Path, issues2: list) -> None:
        self.result_text.config(state="normal")
        self.result_text.delete("1.0", "end")
        self.result_text.insert(
            "end", f"Сохранён исправленный документ:\n{out}\n\n", "ok"
        )
        if not issues2:
            self.result_text.insert("end", "✓ Перепроверка — всё чисто.\n", "ok")
        else:
            self.result_text.insert(
                "end",
                f"После исправления осталось {len(issues2)} замечаний "
                "(их нужно поправить руками):\n\n",
                "info",
            )
            for n, i in enumerate(issues2, 1):
                self.result_text.insert("end", f"  {n}. {i.text}\n", "manual")
        self.result_text.config(state="disabled")
        self.open_fixed_btn.config(state="normal")
        self.fix_btn.config(state="disabled")

    def _open_fixed(self) -> None:
        if self.fixed_path and self.fixed_path.exists():
            try:
                os.startfile(str(self.fixed_path))  # type: ignore[attr-defined]
            except AttributeError:
                import subprocess
                subprocess.Popen(["xdg-open", str(self.fixed_path)])

    def _clear(self) -> None:
        self.current_path = None
        self.issues = []
        self.fixed_path = None
        self.file_label.config(text="(файл не выбран)", foreground="#666")
        self.fix_btn.config(state="disabled")
        self.open_fixed_btn.config(state="disabled")
        self._set_text("Перетащите файл в окно или выберите через кнопку.", tag="info")

    def _set_text(self, msg: str, tag: str = "info") -> None:
        self.result_text.config(state="normal")
        self.result_text.delete("1.0", "end")
        self.result_text.insert("end", msg + "\n", tag)
        self.result_text.config(state="disabled")


def main() -> None:
    if HAS_DND:
        root = TkinterDnD.Tk()
    else:
        root = tk.Tk()
    App(root)
    root.mainloop()


if __name__ == "__main__":
    main()
